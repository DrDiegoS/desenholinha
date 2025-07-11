
import streamlit as st
import openai
from docx import Document
from datetime import datetime
import io

# === CONFIGURAÇÃO OPENROUTER ===
openai.api_key = "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"
openai.api_base = "https://openrouter.ai/api/v1"

st.set_page_config(page_title="Assistente de Linha de Cuidado", layout="centered")
st.title("🩺 Assistente Inteligente de Linha de Cuidado")

# === Função para obter sugestão da IA ===
@st.cache_data(show_spinner=False)
def gerar_sugestao(resposta):
    try:
        completion = openai.ChatCompletion.create(
            model="mistralai/mistral-7b-instruct",
            messages=[
                {"role": "system", "content": "Você é um especialista em gestão clínica e linhas de cuidado."},
                {"role": "user", "content": f"Avalie a seguinte resposta e sugira melhorias: {resposta}"}
            ]
        )
        return completion['choices'][0]['message']['content']
    except Exception as e:
        return f"(Erro ao gerar sugestão: {e})"

# === Formulário ===
with st.form("linha_cuidado_form"):
    nome = st.text_input("📝 Nome da Linha de Cuidado")
    criterio = st.text_area("📌 Critérios de Entrada")
    fases = st.text_area("🔁 Fases ou Etapas Principais")
    atores = st.text_area("👥 Profissionais Envolvidos")
    indicador1 = st.text_input("📊 Indicador 1")
    indicador2 = st.text_input("📊 Indicador 2")
    extras = st.text_area("🧩 Observações / Ferramentas")

    submitted = st.form_submit_button("Gerar Documento")

if submitted:
    s1 = gerar_sugestao(criterio)
    s2 = gerar_sugestao(fases)
    s3 = gerar_sugestao(atores)
    s4 = gerar_sugestao(indicador1)
    s5 = gerar_sugestao(indicador2)
    s6 = gerar_sugestao(extras)

    doc = Document()
    doc.add_heading(f"Linha de Cuidado: {nome}", 0)
    doc.add_heading("1. Critérios de Entrada", level=1)
    doc.add_paragraph(criterio)
    doc.add_heading("2. Fases / Etapas", level=1)
    doc.add_paragraph(fases)
    doc.add_heading("3. Profissionais Envolvidos", level=1)
    doc.add_paragraph(atores)
    doc.add_heading("4. Indicadores", level=1)
    doc.add_paragraph(f"• {indicador1}")
    doc.add_paragraph(f"• {indicador2}")
    doc.add_heading("5. Observações", level=1)
    doc.add_paragraph(extras)
    doc.add_heading("6. Sugestões da IA", level=1)
    doc.add_paragraph("Critérios: " + s1)
    doc.add_paragraph("Fases: " + s2)
    doc.add_paragraph("Atores: " + s3)
    doc.add_paragraph("Indicador 1: " + s4)
    doc.add_paragraph("Indicador 2: " + s5)
    doc.add_paragraph("Extras: " + s6)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("✅ Documento gerado com sucesso!")
    st.download_button("📥 Baixar Documento Word", buffer, file_name=f"Linha_de_Cuidado_{nome.replace(' ', '_')}.docx")
